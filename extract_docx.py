import docx, os, glob, sys
sys.stdout.reconfigure(encoding='utf-8')

folder = r'F:\Kiemtratrunglap\WORD'
files = sorted(glob.glob(os.path.join(folder, '*.docx')))
for f in files:
    doc = docx.Document(f)
    name = os.path.basename(f)
    print(f'========== {name} ==========')
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            print(f'[{i}] {para.text}')
    for ti, table in enumerate(doc.tables):
        print(f'--- Table {ti} ---')
        for ri, row in enumerate(table.rows):
            cells = [c.text.strip() for c in row.cells]
            print(f'  Row {ri}: {" | ".join(cells)}')
    print()
